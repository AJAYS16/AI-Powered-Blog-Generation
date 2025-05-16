from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
import time
import re
from typing import List, Dict
import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai
from dotenv import load_dotenv
from datetime import datetime
import logging
from image_generator import ImageGenerator
from webdriver_manager.chrome import ChromeDriverManager
import random
import docx

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Configure Gemini
genai.configure(api_key=os.getenv('GEMINI_API_KEY'))
model = genai.GenerativeModel('gemini-1.5-pro')

# Define the blog directory path
BLOG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "blogs")

def ensure_blog_directory():
    """Create the blog directory if it doesn't exist"""
    if not os.path.exists(BLOG_DIR):
        os.makedirs(BLOG_DIR)
        logger.info(f"Created blog directory at: {BLOG_DIR}")

class ContentScraper:
    def __init__(self):
        self.setup_driver()
        ensure_blog_directory()
        
    def setup_driver(self):
        """Setup Chrome driver with undetectable Selenium options"""
        try:
            chrome_options = Options()
            
            # Essential Chrome options for stability
            chrome_options.add_argument("--disable-blink-features=AutomationControlled")
            chrome_options.add_argument("--no-sandbox")
            chrome_options.add_argument("--headless=new")  # Run in headless mode
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument("--disable-infobars")
            chrome_options.add_argument("--disable-extensions")
            chrome_options.add_argument("--disable-popup-blocking")
            chrome_options.add_argument("--disable-notifications")
            chrome_options.add_argument("--disable-gpu")
            chrome_options.add_argument("--window-size=1920,1080")
            chrome_options.add_argument("--start-maximized")
            chrome_options.add_argument("--ignore-certificate-errors")
            chrome_options.add_argument(f"user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36")
            
            # Initialize the driver with error handling
            try:
                service = Service(ChromeDriverManager().install())
                self.driver = webdriver.Chrome(service=service, options=chrome_options)
                self.wait = WebDriverWait(self.driver, 10)
                print("Chrome driver initialized successfully")
                
                # Add longer initial wait time after browser launch
                time.sleep(random.uniform(15, 25))  # Wait 15-25 seconds after browser launch
                
                return True
            except Exception as e:
                print(f"Error initializing Chrome driver: {str(e)}")
                return False
            
        except Exception as e:
            print(f"Error setting up Chrome options: {str(e)}")
            return False

    def handle_captcha(self):
        """Handle Google's reCAPTCHA with exponential backoff"""
        try:
            recaptcha_frame = self.driver.find_elements(By.CSS_SELECTOR, "iframe[title*='reCAPTCHA']")
            if recaptcha_frame:
                logger.info("reCAPTCHA detected, implementing backoff strategy...")
                # Exponential backoff
                wait_time = min(300, 2 ** self.captcha_count)  # Max 5 minutes
                logger.info(f"Waiting {wait_time} seconds before retry...")
                time.sleep(wait_time)
                self.captcha_count += 1
                self.driver.refresh()
                return True
                
        except Exception as e:
            logger.warning(f"Error handling CAPTCHA: {str(e)}")
        return False

    def search_google(self, query: str) -> List[str]:
        """Search Google and return top result URLs with detailed logging"""
        try:
            logger.info(f"Starting search for: {query}")
            
            # Add shorter random delay before search
            time.sleep(random.uniform(5, 8))  # Wait 5-8 seconds before starting search
            
            # Navigate to Google with retry logic
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    self.driver.get('https://www.google.com')
                    # Add random delay after page load
                    time.sleep(random.uniform(2, 4))  # Reduced wait time after page load
                    logger.info(f"Successfully navigated to Google (attempt {attempt + 1})")
                    break
                except Exception as e:
                    logger.error(f"Error navigating to Google (attempt {attempt + 1}): {str(e)}")
                    if attempt < max_retries - 1:
                        time.sleep(2)
                        continue
                    return []

            # Wait for and interact with search box
            try:
                search_box = WebDriverWait(self.driver, 20).until(
                    EC.presence_of_element_located((By.NAME, "q"))
                )
                search_box.clear()
                search_box.send_keys(query)
                search_box.send_keys(Keys.RETURN)
                logger.info("Search submitted")
                time.sleep(3)
            except Exception as e:
                logger.error(f"Error with search box: {str(e)}")
                return []

            # Extract URLs with retry logic
            urls = []
            selectors = [
                "div.g div.yuRUbf > a",  # Main result links
                "div.tF2Cxc > div.yuRUbf > a",  # Another format
                "div.g a[href^='http']",  # Any http links in results
                "a[jsname]",  # Links with jsname attribute
                "a[ping]",    # Links with ping attribute
                "a[href^='http']"  # Any http links as fallback
            ]
            
            for selector in selectors:
                try:
                    logger.info(f"Trying selector: {selector}")
                    elements = WebDriverWait(self.driver, 5).until(
                        EC.presence_of_all_elements_located((By.CSS_SELECTOR, selector))
                    )
                    
                    for element in elements[:5]:  # Process first 5 results
                        try:
                            url = element.get_attribute('href')
                            if url and url.startswith('http') and not any(x in url.lower() for x in [
                                'google.com', 'youtube.com', 'facebook.com', 
                                'twitter.com', 'instagram.com', 'linkedin.com'
                            ]):
                                # Try to click the link in a new tab
                                try:
                                    logger.info(f"Opening URL in new tab: {url}")
                                    # Store current window handle
                                    main_window = self.driver.current_window_handle
                                    
                                    # Open link in new tab using JavaScript
                                    self.driver.execute_script("window.open(arguments[0], '_blank');", url)
                                    time.sleep(2)
                                    
                                    # Switch back to main window
                                    self.driver.switch_to.window(main_window)
                                    
                                    urls.append(url)
                                    logger.info(f"Successfully processed URL: {url}")
                                    
                                    if len(urls) >= 5:  # Stop after 5 valid URLs
                                        return urls
                                        
                                except Exception as e:
                                    logger.error(f"Error opening URL: {str(e)}")
                                    urls.append(url)  # Still add URL even if clicking fails
                                    
                        except Exception as e:
                            logger.error(f"Error processing element: {str(e)}")
                            continue
                            
                    if urls:  # If we found URLs with this selector, stop trying others
                        break
                        
                except Exception as e:
                    logger.warning(f"Selector {selector} failed: {str(e)}")
                    continue
            
            logger.info(f"Found {len(urls)} valid URLs")
            return urls
            
        except Exception as e:
            logger.error(f"General error in search: {str(e)}")
            return []
    
    def extract_content(self, url: str) -> Dict[str, str]:
        """Extract content from URL with enhanced error handling and content extraction"""
        try:
            logger.info(f"Extracting content from: {url}")
            
            # Switch to the tab containing our URL
            found_tab = False
            for handle in self.driver.window_handles:
                try:
                    self.driver.switch_to.window(handle)
                    if url in self.driver.current_url:
                        logger.info("Found matching tab for URL")
                        found_tab = True
                        break
                except Exception:
                    continue
            
            if not found_tab:
                logger.warning(f"No tab found for URL {url}, opening new tab")
                # Open in new tab if not found
                self.driver.execute_script("window.open(arguments[0], '_blank');", url)
                time.sleep(2)
                # Switch to the new tab
                self.driver.switch_to.window(self.driver.window_handles[-1])
            
            # Wait for page load
            logger.info("Waiting for page to load...")
            try:
                WebDriverWait(self.driver, 15).until(
                    lambda driver: driver.execute_script('return document.readyState') == 'complete'
                )
                logger.info("Page fully loaded")
                # Add extra time for JavaScript rendering
                time.sleep(3)
            except Exception as e:
                logger.warning(f"Page load timeout: {str(e)}")
            
            # Try to scroll down the page to load lazy content
            try:
                logger.info("Scrolling page to load all content")
                # Scroll down multiple times with pauses
                for i in range(3):
                    self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight / 3 * {});".format(i + 1))
                    time.sleep(1)
                # Scroll back to top
                self.driver.execute_script("window.scrollTo(0, 0);")
            except Exception as e:
                logger.warning(f"Error scrolling page: {str(e)}")
            
            # Get the page source after JavaScript has rendered
            page_source = self.driver.page_source
            logger.info(f"Retrieved page source: {len(page_source)} characters")
            
            # Parse with BeautifulSoup
            soup = BeautifulSoup(page_source, 'html.parser')
            
            # Remove unnecessary elements
            for element in soup.find_all(['script', 'style', 'nav', 'footer', 'header', 'iframe', 'noscript']):
                element.decompose()
            
            # Get title with fallbacks
            title = ''
            try:
                # Try meta title first
                meta_title = soup.find('meta', property='og:title')
                if meta_title:
                    title = meta_title.get('content', '').strip()
                    logger.info(f"Found title from meta: {title}")
                
                # Fallback to regular title
                if not title and soup.title:
                    title = soup.title.string.strip()
                    logger.info(f"Found title from title tag: {title}")
                
                # Fallback to h1
                if not title:
                    h1 = soup.find('h1')
                    if h1:
                        title = h1.text.strip()
                        logger.info(f"Found title from h1: {title}")
            except Exception as e:
                logger.error(f"Error extracting title: {str(e)}")
            
            # Get main content with multiple strategies
            content = ''
            try:
                # Try multiple extraction methods
                extraction_methods = [
                    # Method 1: Look for main content containers
                    lambda: self._extract_from_containers(soup),
                    # Method 2: Extract from all paragraphs
                    lambda: self._extract_from_all_paragraphs(soup),
                    # Method 3: Extract from all divs with substantial text
                    lambda: self._extract_from_text_divs(soup),
                    # Method 4: Extract anything with text as a last resort
                    lambda: self._extract_all_text(soup)
                ]
                
                for method in extraction_methods:
                    extracted_text = method()
                    if extracted_text:
                        content = extracted_text
                        break
                
                logger.info(f"Extracted content length: {len(content)}")
                
                # If no content found, use a simpler approach as fallback
                if not content:
                    logger.warning("No content found with standard methods, using fallback extraction")
                    # Just grab all visible text as a last resort
                    all_text = self.driver.find_element(By.TAG_NAME, "body").text
                    if len(all_text) > 100:  # Only use if substantial
                        content = all_text
                        logger.info(f"Extracted {len(content)} characters from direct browser text")
                
            except Exception as e:
                logger.error(f"Error extracting content: {str(e)}")
            
            # Clean up the content
            if content:
                # Remove excessive whitespace
                content = re.sub(r'\s+', ' ', content).strip()
                # Remove any remaining HTML
                content = re.sub(r'<[^>]+>', '', content)
                # Fix common encoding issues
                content = content.replace('â€™', "'").replace('â€"', '-').replace('â€œ', '"').replace('â€', '"')
                logger.info(f"Final content length after cleaning: {len(content)}")
            
            # Use title as part of the content if content is too short
            if title and (not content or len(content) < 100):
                content = title + "\n\n" + content
                logger.info("Added title to content due to short content length")
            
            # Close the current tab if it's not the main window
            if len(self.driver.window_handles) > 1:
                self.driver.close()
                # Switch back to the main window
                self.driver.switch_to.window(self.driver.window_handles[0])
            
            return {
                'url': url,
                'title': title,
                'content': content
            }
            
        except Exception as e:
            logger.error(f"General error extracting content from {url}: {str(e)}")
            # Make sure we're back on the main window
            try:
                if len(self.driver.window_handles) > 1:
                    self.driver.close()
                self.driver.switch_to.window(self.driver.window_handles[0])
            except:
                pass
            return {'url': url, 'title': '', 'content': ''}
            
    def _extract_from_containers(self, soup):
        """Extract content from main content containers"""
        logger.info("Extracting from main content containers")
        content = ""
        
        # Try common content containers
        containers = []
        for selector in ['main', 'article', 'div[class*="content"]', 'div[class*="article"]', 'div[class*="post"]', 'div[id*="content"]', 'div[id*="article"]', '.blog-content', '.post-content']:
            try:
                if selector.startswith('.'):
                    elements = soup.select(selector[1:])
                else:
                    elements = soup.find_all(selector.split('[')[0])
                containers.extend(elements)
            except:
                continue
        
        for container in containers:
            try:
                for element in container.find_all(['p', 'li', 'h2', 'h3', 'h4']):
                    text = element.text.strip()
                    if len(text) > 20:  # Only take substantial paragraphs
                        content += text + '\n\n'
            except:
                continue
                
        if content:
            logger.info(f"Extracted {len(content)} characters from content containers")
        
        return content
        
    def _extract_from_all_paragraphs(self, soup):
        """Extract content from all paragraphs"""
        logger.info("Extracting from all paragraphs")
        content = ""
        
        try:
            for p in soup.find_all('p'):
                text = p.text.strip()
                if len(text) > 20:
                    content += text + '\n\n'
        except Exception as e:
            logger.warning(f"Error extracting from paragraphs: {str(e)}")
            
        if content:
            logger.info(f"Extracted {len(content)} characters from paragraphs")
            
        return content
        
    def _extract_from_text_divs(self, soup):
        """Extract content from divs with substantial text"""
        logger.info("Extracting from text divs")
        content = ""
        
        try:
            for div in soup.find_all('div'):
                # Skip divs with many nested divs (likely not content)
                if len(div.find_all('div')) > 5:
                    continue
                    
                text = div.text.strip()
                if len(text) > 100 and len(text.split()) > 20:
                    content += text + '\n\n'
                    # Avoid duplicating content
                    if len(content) > 2000:
                        break
        except Exception as e:
            logger.warning(f"Error extracting from text divs: {str(e)}")
            
        if content:
            logger.info(f"Extracted {len(content)} characters from text divs")
            
        return content
        
    def _extract_all_text(self, soup):
        """Extract all text from the page as a last resort"""
        logger.info("Extracting all text (last resort)")
        content = ""
        
        try:
            # Get text from body
            body = soup.find('body')
            if body:
                content = body.text.strip()
                
                # Simple cleanup to remove menu items and other non-content
                lines = content.split('\n')
                filtered_lines = []
                for line in lines:
                    line = line.strip()
                    # Keep only substantial lines
                    if len(line) > 20:
                        filtered_lines.append(line)
                
                content = '\n\n'.join(filtered_lines)
        except Exception as e:
            logger.warning(f"Error extracting all text: {str(e)}")
            
        if content:
            logger.info(f"Extracted {len(content)} characters from all text")
            
        return content

    def generate_blog_content(self, topic: str) -> List[Dict[str, str]]:
        """Generate blog content by searching and extracting information"""
        logger.info(f"Starting blog content generation for topic: {topic}")
        
        # Search for URLs
        logger.info("Searching Google for relevant URLs")
        urls = self.search_google(topic)
        
        if not urls:
            logger.warning("No URLs found for the topic")
            return []
            
        logger.info(f"Found {len(urls)} URLs to extract content from")
        
        # Extract content from each URL
        contents = []
        for i, url in enumerate(urls):
            logger.info(f"Processing URL {i+1}/{len(urls)}: {url}")
            try:
                content = self.extract_content(url)
                if content['content']:
                    logger.info(f"Successfully extracted content from URL {i+1}")
                    contents.append(content)
                else:
                    logger.warning(f"No content extracted from URL {i+1}")
            except Exception as e:
                logger.error(f"Error processing URL {i+1}: {str(e)}")
                continue
        
        logger.info(f"Successfully extracted content from {len(contents)} URLs")
        return contents
        
    def close(self):
        """Safely close the browser"""
        try:
            if hasattr(self, 'driver'):
                # Instead of quit(), which closes the entire browser, use close() to close only the automated tab
                self.driver.close()
                logger.info("Chrome driver tab closed successfully")
        except Exception as e:
            logger.error(f"Error closing Chrome driver tab: {str(e)}")

def generate_blog(topic: str, contents: List[Dict[str, str]], max_retries=3, style_prompt=None):
    """Generate a comprehensive blog post using Gemini"""
    for attempt in range(max_retries):
        try:
            logger.info(f"Generating blog for topic: {topic} (attempt {attempt + 1})")
            
            # Create a focused prompt based on the collected content
            content_insights = "\n".join([
                f"Source {i+1} ({content['url']}):\n{content['content'][:1000]}..."
                for i, content in enumerate(contents)
            ])
            
            # Use custom style prompt if provided, otherwise use the default prompt
            if style_prompt:
                # Use the custom style prompt
                combined_prompt = f"""You are an expert technology analyst and AI researcher specializing in creating in-depth blog posts about cutting-edge technology.

{style_prompt}

Use the following research material:

{content_insights}

Follow this structure:
1. Create an engaging title
2. Write an introduction establishing context
3. Cover technical aspects and key features
4. Discuss current limitations and challenges
5. Explore applications and impacts
6. Mention future implications
7. Provide a conclusion

Remember to:
- Back claims with specific data points
- Provide concrete examples
- Address both benefits and limitations
"""
            else:
                # Use the original detailed prompt
                combined_prompt = f"""You are an expert technology analyst and AI researcher specializing in creating in-depth, authoritative blog posts about cutting-edge technology.
Your writing combines technical expertise with clear explanations, making complex topics accessible while maintaining professional depth.
Focus on providing comprehensive analysis, technical insights, and industry implications while backing claims with specific examples and data.

Create a professional, in-depth technical blog post about {topic} using the following research material:

{content_insights}

Follow this precise structure and guidelines:

# [Create an SEO-optimized, compelling title that emphasizes technological advancement]

[Initial 2-3 sentence introduction establishing context and significance]

This blog will explore {topic}'s key features, technical advancements, current limitations, and the new possibilities it opens for various industries.

## Technical Deep Dive: What Makes {topic} Different?
[Write 500-700 words that:
- Explain core technical innovations and architectural improvements
- Compare with previous/existing technologies
- Highlight specific technical advantages using bullet points or numbered lists
- Include relevant technical specifications and metrics]

## Overcoming Previous Limitations
[Write 500-700 words covering:
1. Identify 3-4 major limitations of existing solutions
2. Explain how {topic} addresses each limitation
3. Provide specific examples and use cases
4. Include technical details of improvements]

## Current Limitations and Challenges
[Write 500-700 words analyzing:
1. Technical constraints
2. Implementation challenges
3. Resource requirements
4. Performance trade-offs]

## Industry Applications and Impact
[Write 400-500 words examining:
1. Specific industry use cases
2. Implementation examples
3. Business benefits and ROI potential
4. Integration considerations]

## Future Implications and Development
[Write 300-400 words covering:
1. Development roadmap
2. Potential improvements
3. Industry trends
4. Future research directions]

## Technical Specifications and Requirements
[Write 300-400 words detailing:
1. System requirements
2. Performance metrics
3. Integration prerequisites
4. Technical best practices]

## Conclusion
[Write 200-250 words that:
1. Summarize key technical advantages
2. Highlight industry impact
3. Provide forward-looking perspective
4. End with a powerful statement about future potential]

Style Requirements:
1. Maintain professional, authoritative tone throughout
2. Use technical terminology accurately and consistently
3. Include specific metrics, benchmarks, and performance data
4. Break down complex concepts systematically
5. Use bullet points and numbered lists for better readability
6. Include relevant technical diagrams or flowcharts descriptions
7. Cite specific examples and use cases
8. Maintain logical flow between sections

Focus Areas:
- Technical accuracy and depth
- Industry-specific implications
- Implementation considerations
- Performance metrics and benchmarks
- Future development potential
- Practical applications

Remember to:
- Back claims with specific data points
- Provide concrete examples
- Include technical specifications
- Address both benefits and limitations
- Maintain professional language throughout
- Focus on enterprise and industry relevance"""

            # Generate blog using Gemini
            logger.info("Sending prompt to Gemini")
            response = model.generate_content(
                combined_prompt,
                generation_config={
                    "temperature": 0.7,
                    "top_p": 1,
                    "top_k": 40,
                    "max_output_tokens": 8192,
                }
            )
            
            # Extract the generated content
            blog_content = response.text
            logger.info(f"Generated blog content: {len(blog_content)} characters")
            
            return {
                'topic': topic,
                'content': blog_content,
                'sources': [content['url'] for content in contents],
                'timestamp': datetime.now().strftime("%Y-%m-%d_%H-%M-%S"),
                'model': 'gemini-1.5-pro'
            }
            
        except Exception as e:
            logger.error(f"Error generating blog (attempt {attempt + 1}): {str(e)}")
            if attempt < max_retries - 1:
                time.sleep(5)
            else:
                raise
                
def save_blog_to_word(blog: Dict):
    """Save generated blog to a Word document with generated images"""
    try:
        doc = Document()
        
        # Add title
        title = blog['content'].split('\n')[0].replace('#', '').strip()
        doc.add_heading(title, 0)
        
        # Split content into lines for processing
        lines = blog['content'].split('\n')
        
        # Find where the main content starts (after title)
        start_idx = 1
        while start_idx < len(lines) and not lines[start_idx].strip():
            start_idx += 1
            
        # Check if we have social media content to include
        if 'social_content' in blog and blog['social_content']:
            # Add social media posts section
            has_content = False
            
            for platform, posts in blog['social_content'].items():
                if not posts:
                    continue
                    
                if not has_content:
                    # Add section header only once
                    doc.add_heading("Recent Social Media Updates", level=2)
                    has_content = True
                
                # Add platform-specific heading
                platform_name = platform.capitalize()
                if platform == "mock":
                    platform_name = "Social Media"
                    
                doc.add_heading(f"Recent {platform_name} Posts", level=3)
                
                # Add each post
                for post in posts[:3]:  # Limit to 3 posts per platform
                    # Add post author
                    author_para = doc.add_paragraph(style='Intense Quote')
                    author_para.add_run(post['author']).bold = True
                    
                    # Add post text
                    doc.add_paragraph(post['text'])
                    
                    # Add date if available
                    if post.get('date'):
                        try:
                            date_obj = datetime.fromisoformat(post['date'].replace('Z', '+00:00'))
                            formatted_date = date_obj.strftime("%b %d, %Y")
                            date_para = doc.add_paragraph()
                            date_para.add_run(formatted_date).italic = True
                        except:
                            date_para = doc.add_paragraph()
                            date_para.add_run(post['date']).italic = True
                    
                    # Add URL if available
                    if post.get('url'):
                        url_para = doc.add_paragraph()
                        url_para.add_run("View Original Post").italic = True
                        
                        # Add hyperlink
                        add_hyperlink(url_para, post['url'], 'View Original Post', 'Hyperlink')
                    
                    # Add separator between posts
                    doc.add_paragraph('---')
            
            # Add a page break after the social media section if we added content
            if has_content:
                doc.add_page_break()
        
        # Process content with or without images
        if 'generated_images' not in blog or not blog['generated_images']:
            # Simply add all content as text
            for line in lines[start_idx:]:
                if line.startswith('##'):
                    doc.add_heading(line.replace('##', '').strip(), 1)
                elif line.startswith('#'):
                    doc.add_heading(line.replace('#', '').strip(), 0)
                else:
                    doc.add_paragraph(line)
        else:
            # Add content with images at appropriate sections
            image_counter = 0
            has_added_first_image = False
            
            # Add first image near the beginning if available
            if len(blog['generated_images']) > 0 and not has_added_first_image:
                try:
                    # Add first image after first paragraph
                    first_para_idx = next((i for i in range(start_idx, min(start_idx + 10, len(lines))) 
                                         if lines[i].strip() and not lines[i].startswith('#')), None)
                    
                    if first_para_idx:
                        # Add paragraphs up to first image
                        for line in lines[start_idx:first_para_idx+1]:
                            if line.startswith('##'):
                                doc.add_heading(line.replace('##', '').strip(), 1)
                            elif line.startswith('#'):
                                doc.add_heading(line.replace('#', '').strip(), 0)
                            else:
                                doc.add_paragraph(line)
                        
                        # Add first image
                        doc.add_picture(blog['generated_images'][0], width=Inches(6))
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        image_counter = 1
                        has_added_first_image = True
                        start_idx = first_para_idx + 1
                except Exception as e:
                    logger.warning(f"Failed to add first image: {str(e)}")
            
            # Process remaining content and add images before major sections
            for i, line in enumerate(lines[start_idx:], start=start_idx):
                if line.startswith('##'):
                    # Add image before each major section if available
                    if image_counter < len(blog['generated_images']):
                        try:
                            doc.add_picture(blog['generated_images'][image_counter], width=Inches(6))
                            paragraph = doc.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            image_counter += 1
                        except Exception as e:
                            logger.warning(f"Failed to add image {image_counter}: {str(e)}")
                
                    doc.add_heading(line.replace('##', '').strip(), 1)
                elif line.startswith('#'):
                    doc.add_heading(line.replace('#', '').strip(), 0)
                else:
                    doc.add_paragraph(line)
            
            # Add any remaining images at the end
            while image_counter < len(blog['generated_images']):
                try:
                    doc.add_picture(blog['generated_images'][image_counter], width=Inches(6))
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    image_counter += 1
                except Exception as e:
                    logger.warning(f"Failed to add remaining image {image_counter}: {str(e)}")
        
        # Create a unique filename with date and part number
        base_filename = blog['topic'].replace(' ', '_').replace('-', '_')
        filename = f"{base_filename}_{blog.get('timestamp', datetime.now().strftime('%Y-%m-%d_%H-%M-%S'))}.docx"
        filepath = os.path.join(BLOG_DIR, filename)
        doc.save(filepath)
        logger.info(f"Blog saved to: {filepath}")
        return filepath
        
    except Exception as e:
        logger.error(f"Error saving blog to Word: {str(e)}")
        raise

def add_hyperlink(paragraph, url, text, style):
    """Add a hyperlink to a paragraph"""
    # This gets access to the document.xml.rels file and gets a new relation id
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the hyperlink and add it to the paragraph
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    # Create a run that provides the actual text that becomes a hyperlink
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add/Set the style of the text in the hyperlink
    if style:
        rStyle = docx.oxml.shared.OxmlElement('w:rStyle')
        rStyle.set(docx.oxml.shared.qn('w:val'), style)
        rPr.append(rStyle)

    # Add the hyperlink properties (underline, blue text)
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)
    
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000FF')
    rPr.append(c)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    
    return hyperlink

def main():
    scraper = None
    try:
        # Get topic and number of days from user
        topic = input("Enter the topic to research and create blogs about: ")
        logger.info(f"User entered topic: {topic}")
        
        while True:
            try:
                num_days = int(input("Enter the number of days (blogs) needed: "))
                if num_days > 0:
                    logger.info(f"User requested {num_days} blogs")
                    break
                print("Please enter a positive number.")
            except ValueError:
                print("Please enter a valid number.")
        
        # Initialize scraper
        logger.info("Initializing content scraper")
        scraper = ContentScraper()
        logger.info(f"Generating {num_days} blogs about: {topic}")
        
        # Generate multiple blogs
        for day in range(1, num_days + 1):
            logger.info(f"Starting blog generation {day}/{num_days}")
            print(f"\nGenerating blog {day} of {num_days}...")
            
            # Modify search query to get diverse results for each day
            search_queries = [
                f"{topic} ",
            ]
            
            all_contents = []
            for query in search_queries:
                logger.info(f"Searching for content with query: {query}")
                try:
                    contents = scraper.generate_blog_content(query)
                    # Take only first 2 URLs for each query
                    all_contents.extend(contents[:2])
                    logger.info(f"Found {len(contents)} contents for query: {query}")
                except Exception as e:
                    logger.error(f"Error generating blog content for query '{query}': {str(e)}")
                    continue
                
                # Add delay between searches to avoid detection
                logger.info("Adding delay between searches")
                time.sleep(3)
            
            if not all_contents:
                logger.warning(f"No relevant content found for blog {day}")
                print(f"No relevant content found for blog {day}. Skipping...")
                continue
            
            # Use all collected contents (should be up to 6 sources - 2 from each query)
            relevant_contents = all_contents
            logger.info(f"Found {len(relevant_contents)} relevant sources for blog {day}")
            print(f"Found {len(relevant_contents)} sources for blog {day}")
            
            # Generate blog using Gemini
            logger.info(f"Generating content for blog {day}")
            print(f"Generating content for blog {day}...")
            try:
                blog = generate_blog(f"{topic} - Part {day}", relevant_contents)
                logger.info("Blog content generated successfully")
            except Exception as e:
                logger.error(f"Error generating blog content: {str(e)}")
                print(f"Error generating blog content: {str(e)}")
                continue
            
            # Save blog to Word document
            logger.info(f"Saving blog {day} to Word document")
            print(f"Saving blog {day}...")
            try:
                filepath = save_blog_to_word(blog)
                logger.info(f"Blog {day} saved to: {filepath}")
                print(f"Blog {day} saved to: {filepath}")
            except Exception as e:
                logger.error(f"Error saving blog to Word: {str(e)}")
                print(f"Error saving blog to Word: {str(e)}")
                continue
            
            # Add delay between blog generations
            if day < num_days:
                delay = 10  # seconds
                logger.info(f"Waiting {delay} seconds before generating next blog")
                print(f"Waiting {delay} seconds before generating next blog...")
                time.sleep(delay)
        
        logger.info("All blogs generated successfully")
        print("\nAll blogs generated successfully!")
        
    except Exception as e:
        logger.error(f"An error occurred: {str(e)}")
        print(f"An error occurred: {str(e)}")
    finally:
        if scraper:
            logger.info("Closing scraper")
            scraper.close()

if __name__ == "__main__":
    main()
